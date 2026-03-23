#!/usr/bin/env python3
"""
Word文档双引号替换工具
将所有英文双引号替换为中文双引号，不修改其他样式。
用法: python replace-double-quotes.py document.docx
"""

import argparse
import re
import shutil
from pathlib import Path

from docx import Document


def replace_quotes_in_text(text: str) -> str:
    """将文本中的英文双引号替换为中文双引号（成对匹配）"""
    result = []
    open_quote = True
    for ch in text:
        if ch == '"':
            result.append('\u201c' if open_quote else '\u201d')
            open_quote = not open_quote
        else:
            result.append(ch)
    return ''.join(result)


def process_runs(runs):
    """处理一组run中的双引号替换"""
    for run in runs:
        if '"' in run.text:
            run.text = replace_quotes_in_text(run.text)


def replace_quotes_in_document(input_path: str):
    """替换文档中所有双引号"""
    document = Document(input_path)

    # 处理正文段落
    for paragraph in document.paragraphs:
        process_runs(paragraph.runs)

    # 处理表格
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_runs(paragraph.runs)

    # 处理页眉页脚
    for section in document.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for paragraph in header.paragraphs:
                    process_runs(paragraph.runs)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for paragraph in footer.paragraphs:
                    process_runs(paragraph.runs)

    return document


def main():
    parser = argparse.ArgumentParser(description='Word文档双引号替换工具')
    parser.add_argument('input', help='输入的Word文档路径')

    args = parser.parse_args()
    input_path = Path(args.input)

    if not input_path.exists():
        print(f"错误: 输入文件不存在: {input_path}")
        return 1

    # 备份原文件
    backup_path = input_path.parent / f"{input_path.stem}_backup{input_path.suffix}"
    shutil.copy2(input_path, backup_path)
    print(f"原文件已备份: {backup_path}")

    # 替换并保存
    document = replace_quotes_in_document(str(input_path))
    document.save(input_path)
    print(f"双引号替换完成: {input_path}")

    return 0


if __name__ == '__main__':
    exit(main())
