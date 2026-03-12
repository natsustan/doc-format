#!/usr/bin/env python3
"""
Word文档格式化工具
用法: python format_doc.py document.docx [--config config.yaml]
"""

import argparse
import os
import re
import shutil
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def load_config(config_path: str) -> dict:
    """加载配置文件"""
    with open(config_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)


def get_align(align_str: str) -> WD_ALIGN_PARAGRAPH:
    """字符串转对齐枚举"""
    mapping = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def set_run_font(run, font_name: str, font_size: float = None):
    """设置run的字体，保留原有加粗/斜体等属性"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size is not None:
        run.font.size = Pt(font_size)


def set_paragraph_spacing(paragraph, line_spacing: float):
    """设置段落行间距"""
    paragraph.paragraph_format.line_spacing = line_spacing


def set_first_line_indent(paragraph, chars: int, font_size: float):
    """设置首行缩进（按字符数）"""
    indent = Pt(font_size * chars)
    paragraph.paragraph_format.first_line_indent = indent


def get_style_name(paragraph) -> str:
    """获取段落样式名称（标准化）"""
    if paragraph.style is None:
        return 'normal'
    style_name = paragraph.style.name.lower().replace(' ', '')
    return style_name


def is_empty_paragraph(paragraph) -> bool:
    """判断段落是否为空"""
    text = paragraph.text.strip()
    return len(text) == 0 and len(paragraph.runs) == 0


def has_image(paragraph) -> bool:
    """判断段落是否包含图片"""
    for run in paragraph.runs:
        if run._element.xpath('.//a:blip'):
            return True
        if run._element.xpath('.//w:drawing'):
            return True
    return False


def format_paragraph(paragraph, config: dict):
    """格式化段落"""
    style_name = get_style_name(paragraph)
    styles_config = config['styles']
    font_name = config['font_name']
    line_spacing = config['line_spacing']
    
    style_cfg = None
    if style_name == 'title':
        style_cfg = styles_config.get('title')
    elif style_name == 'heading1':
        style_cfg = styles_config.get('heading1')
    elif style_name == 'heading2':
        style_cfg = styles_config.get('heading2')
    elif style_name == 'heading3':
        style_cfg = styles_config.get('heading3')
    elif style_name == 'heading4':
        style_cfg = styles_config.get('heading4')
    else:
        style_cfg = styles_config.get('normal')
    
    if style_cfg is None:
        style_cfg = styles_config.get('normal', {'size': 12, 'align': 'left'})
    
    font_size = style_cfg.get('size', 12)
    align = style_cfg.get('align', 'left')
    first_indent = style_cfg.get('first_indent', 0)
    
    paragraph.alignment = get_align(align)
    set_paragraph_spacing(paragraph, line_spacing)
    
    if first_indent > 0 and style_name not in ['title', 'heading1', 'heading2', 'heading3', 'heading4']:
        set_first_line_indent(paragraph, first_indent, font_size)
    
    for run in paragraph.runs:
        set_run_font(run, font_name, font_size)


def format_image_paragraph(paragraph):
    """图片段落居中"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def table_has_borders(table) -> bool:
    """判断表格是否已设置边框"""
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        return False
    
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is None:
        return False
    
    for border in tbl_borders:
        if border.get(qn('w:val')) not in (None, 'nil', 'none'):
            return True
    
    return False


def set_table_borders(table):
    """为表格设置默认边框"""
    tbl_pr = table._element.tblPr
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tbl_borders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tbl_borders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')


def format_table(table, config: dict):
    """格式化表格"""
    table_cfg = config.get('table', {})
    font_name = table_cfg.get('font_name', config['font_name'])
    font_size = table_cfg.get('font_size', 10)
    
    if not table_has_borders(table):
        set_table_borders(table)
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text
                stripped = text.strip()
                if text != stripped:
                    for run in paragraph.runs:
                        run.text = run.text.strip()
                
                for run in paragraph.runs:
                    set_run_font(run, font_name, font_size)


def format_list_paragraph(paragraph, config: dict):
    """格式化列表段落，统一使用小圆点"""
    bullet_char = config.get('list', {}).get('bullet_char', '•')
    
    numPr = paragraph._element.pPr.numPr if paragraph._element.pPr is not None else None
    if numPr is not None:
        pass


def set_page_margins(document, config: dict):
    """设置页边距"""
    margin_cfg = config.get('page_margin', {})
    for section in document.sections:
        section.top_margin = Cm(margin_cfg.get('top', 2.54))
        section.bottom_margin = Cm(margin_cfg.get('bottom', 2.54))
        section.left_margin = Cm(margin_cfg.get('left', 3.17))
        section.right_margin = Cm(margin_cfg.get('right', 3.17))


def remove_extra_empty_lines(document):
    """删除多余的空行，连续空行只保留一个"""
    paragraphs = document.paragraphs
    to_remove = []
    prev_empty = False
    
    for i, para in enumerate(paragraphs):
        is_empty = is_empty_paragraph(para) and not has_image(para)
        if is_empty and prev_empty:
            to_remove.append(para)
        prev_empty = is_empty
    
    for para in to_remove:
        p = para._element
        p.getparent().remove(p)


def is_list_paragraph(paragraph) -> bool:
    """判断是否为列表段落"""
    if paragraph._element.pPr is None:
        return False
    numPr = paragraph._element.pPr.find(qn('w:numPr'))
    return numPr is not None


def convert_to_bullet_list(paragraph, bullet_char: str = '•'):
    """将列表转换为使用指定符号的无序列表样式"""
    pass


def format_document(input_path: str, config: dict):
    """格式化整个文档"""
    document = Document(input_path)
    
    set_page_margins(document, config)
    
    for paragraph in document.paragraphs:
        if has_image(paragraph):
            format_image_paragraph(paragraph)
            continue
        
        if is_list_paragraph(paragraph):
            format_paragraph(paragraph, config)
            continue
        
        format_paragraph(paragraph, config)
    
    for table in document.tables:
        format_table(table, config)
    
    remove_extra_empty_lines(document)
    
    return document


def backup_and_save(input_path: str, document):
    """备份原文件并保存新文件"""
    input_path = Path(input_path)
    backup_path = input_path.parent / f"{input_path.stem}_backup{input_path.suffix}"
    
    shutil.copy2(input_path, backup_path)
    print(f"原文件已备份: {backup_path}")
    
    document.save(input_path)
    print(f"格式化完成: {input_path}")


def main():
    parser = argparse.ArgumentParser(description='Word文档格式化工具')
    parser.add_argument('input', help='输入的Word文档路径')
    parser.add_argument('--config', '-c', default=None, help='配置文件路径（默认使用同目录下的config.yaml）')
    
    args = parser.parse_args()
    
    if args.config:
        config_path = args.config
    else:
        script_dir = Path(__file__).parent
        config_path = script_dir / 'config.yaml'
    
    if not os.path.exists(config_path):
        print(f"错误: 配置文件不存在: {config_path}")
        return 1
    
    if not os.path.exists(args.input):
        print(f"错误: 输入文件不存在: {args.input}")
        return 1
    
    config = load_config(config_path)
    document = format_document(args.input, config)
    backup_and_save(args.input, document)
    
    return 0


if __name__ == '__main__':
    exit(main())
